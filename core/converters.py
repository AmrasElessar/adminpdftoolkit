"""Conversion utilities (Section B) — pure-Python wheels only.

Section B fills the "non-PDF input → PDF" gap and the reverse "PDF → tabular".
Every helper sticks to the all-in-one constraint: pip-installable wheels with
no system dependency. Trade-offs vs LibreOffice/headless-Chrome paths:
  • docx → PDF emits plain HTML and pipes through xhtml2pdf — text + basic
    formatting are preserved; SmartArt / complex layouts lose fidelity.
  • xlsx → PDF renders one HTML table per sheet — formulas are evaluated to
    their cached value, charts are dropped.
  • HTML → PDF supports the xhtml2pdf CSS subset (no flexbox, no JS, no
    remote stylesheet fetching).
"""

from __future__ import annotations

from pathlib import Path
from typing import Any

from .errors import sanitize_error
from .logging_setup import logger
from .pdf_tools import _find_unicode_font, _save_pdf

# ----- Image → PDF --------------------------------------------------------
_IMAGE_EXTS: frozenset[str] = frozenset(
    {".jpg", ".jpeg", ".png", ".webp", ".bmp", ".tif", ".tiff", ".gif"}
)


def image_to_pdf(images: list[Path], output: Path) -> int:
    """Bundle a list of images into a single PDF, one image per page.

    Returns the resulting page count. Each image is rendered at its native
    pixel size; the page dimensions are derived from the image so no rescale
    happens. Supported formats: jpg / jpeg / png / webp / bmp / tif / tiff /
    gif (PyMuPDF + Pillow handle decoding).
    """
    import fitz

    if not images:
        raise ValueError("Dönüştürülecek görsel yok.")
    for p in images:
        if p.suffix.lower() not in _IMAGE_EXTS:
            raise ValueError(f"Desteklenmeyen görsel formatı: {p.name}")

    out_doc = fitz.open()
    try:
        for img_path in images:
            try:
                with fitz.open(str(img_path)) as img_doc:
                    pdf_bytes = img_doc.convert_to_pdf()
                with fitz.open(stream=pdf_bytes, filetype="pdf") as tmp:
                    out_doc.insert_pdf(tmp)
            except Exception as e:
                raise ValueError(f"{img_path.name} okunamadı: {sanitize_error(e)}") from e
        page_count = int(out_doc.page_count)
        _save_pdf(out_doc, output)
    finally:
        out_doc.close()
    return page_count


# ----- PDF → Markdown -----------------------------------------------------
def pdf_to_markdown(input_path: Path, output: Path) -> int:
    """Extract text + heading structure from a PDF and write Markdown.

    Two passes:
      1. Collect every span's font-size to derive heading thresholds
         (median × 1.15 / 1.3 / 1.6 → H3 / H2 / H1).
      2. Emit per-line text, prefixing with ``#``/``##``/``###`` when the
         line's max font size crosses a threshold; bold-flagged runs are
         wrapped in ``**…**`` (only when the line itself isn't a heading).

    Refuses early on image-only PDFs (uses ``classify_pdf_extractability``)
    so the user gets a clear message instead of an empty .md file.
    """
    import core

    extractability = core.classify_pdf_extractability(input_path)
    if not extractability["extractable"]:
        raise ValueError(extractability["message"])

    import fitz

    with fitz.open(str(input_path)) as doc:
        if doc.is_encrypted and not doc.authenticate(""):
            raise ValueError("Şifreli PDF — önce şifreyi kaldırın.")
        page_count = int(doc.page_count)

        # Pass 1: collect all span sizes to derive heading thresholds
        sizes: list[float] = []
        for page in doc:
            for block in page.get_text("dict").get("blocks", []):
                if block.get("type") != 0:
                    continue
                for line in block.get("lines", []):
                    for span in line.get("spans", []):
                        s = float(span.get("size", 0) or 0)
                        if s > 0:
                            sizes.append(s)
        if not sizes:
            output.parent.mkdir(parents=True, exist_ok=True)
            output.write_text("", encoding="utf-8")
            return page_count
        sizes.sort()
        median = sizes[len(sizes) // 2]
        h1_cut = median * 1.6
        h2_cut = median * 1.3
        h3_cut = median * 1.15

        # Pass 2: emit
        lines_out: list[str] = []
        for pidx, page in enumerate(doc):
            if pidx > 0:
                lines_out.append("")
                lines_out.append(f"<!-- Sayfa {pidx + 1} -->")
                lines_out.append("")
            for block in page.get_text("dict").get("blocks", []):
                if block.get("type") != 0:
                    continue
                for line in block.get("lines", []):
                    parts: list[str] = []
                    max_size = 0.0
                    for span in line.get("spans", []):
                        text = (span.get("text") or "").strip()
                        if not text:
                            continue
                        sz = float(span.get("size", 0) or 0)
                        if sz > max_size:
                            max_size = sz
                        flags = int(span.get("flags", 0))
                        is_bold = bool(flags & 16)  # 0x10 in PyMuPDF
                        if is_bold and max_size < h3_cut:
                            text = f"**{text}**"
                        parts.append(text)
                    if not parts:
                        continue
                    line_text = " ".join(parts)
                    if max_size >= h1_cut:
                        lines_out.append(f"# {line_text}")
                    elif max_size >= h2_cut:
                        lines_out.append(f"## {line_text}")
                    elif max_size >= h3_cut:
                        lines_out.append(f"### {line_text}")
                    else:
                        lines_out.append(line_text)
                lines_out.append("")  # paragraph break between blocks

    output.parent.mkdir(parents=True, exist_ok=True)
    cleaned: list[str] = []
    blank = False
    for ln in lines_out:
        if not ln.strip():
            if not blank:
                cleaned.append("")
            blank = True
        else:
            cleaned.append(ln)
            blank = False
    body = "\n".join(cleaned).strip() + "\n"
    # Cap markdown output at 50 MB. A pathological PDF (millions of tiny
    # text spans) could produce a ridiculously large .md otherwise.
    max_bytes = 50 * 1024 * 1024
    encoded = body.encode("utf-8")
    if len(encoded) > max_bytes:
        encoded = encoded[:max_bytes]
        # Trim to a clean UTF-8 boundary
        body = encoded.decode("utf-8", errors="ignore")
        body += "\n\n<!-- Çıktı 50 MB üst sınırına ulaştığı için kesildi. -->\n"
    output.write_text(body, encoding="utf-8")
    return page_count


# ----- PDF → CSV ----------------------------------------------------------
def pdf_to_csv(
    input_path: Path,
    output: Path,
    *,
    table_index: int | None = None,
    delimiter: str = ",",
) -> int:
    """Extract embedded tables from a PDF and write a CSV. Delegates to parsers.

    ``table_index`` (1-indexed) picks one specific table; ``None`` concatenates
    every table found, separating each with a blank row. Returns the number of
    rows written (excluding separators).
    """
    from parsers.generic_table import GenericTableParser

    return GenericTableParser().to_csv(
        input_path,
        output,
        table_index=table_index,
        delimiter=delimiter,
    )


# ----- HTML / URL → PDF ---------------------------------------------------
_PISA_FONT_FAMILY: list[str] = []


def _sibling_font(regular_path: str, variant: str) -> str | None:
    """Look up a bold/italic sibling alongside ``regular_path``."""
    p = Path(regular_path)
    suffix = p.suffix
    table = {
        "bold": ("bd", "-bold", "b"),
        "italic": ("i", "-italic", "it"),
        "bolditalic": ("bi", "-boldoblique", "z", "-bolditalic"),
    }
    candidates: list[Path] = []
    for tag in table.get(variant, ()):
        candidates.append(p.with_name(p.stem + tag + suffix))
        for marker in ("-regular", "-Regular", "regular", "Regular"):
            if marker in p.stem:
                replacement = {
                    "bold": "Bold",
                    "italic": "Italic",
                    "bolditalic": "BoldItalic",
                }[variant]
                candidates.append(p.with_name(p.stem.replace(marker, replacement) + suffix))
    for c in candidates:
        try:
            if c.exists() and c.is_file():
                return str(c)
        except OSError:
            continue
    return None


def _ensure_pisa_unicode_font() -> str:
    """Resolve a Unicode font family name for xhtml2pdf.

    xhtml2pdf's built-in fonts are Latin-1 only — Turkish glyphs are
    silently substituted to ASCII look-alikes. To get correct rendering we
    must register the TTF *through xhtml2pdf's own context*, because
    xhtml2pdf maintains a parallel ``fontList`` lookup separate from
    reportlab's ``pdfmetrics``. We perform the registration via a
    ``<style>@font-face …</style>`` block whose ``url(...)`` resolves through
    a ``link_callback`` (see ``_xhtml2pdf_render``). This function only
    decides whether a font is available and returns its family name; the
    actual loading happens per-render through the callback.
    Returns ``"HTUni"`` if a system TTF was found, else ``"Helvetica"``.
    """
    if _PISA_FONT_FAMILY:
        return _PISA_FONT_FAMILY[0]
    font_path = _find_unicode_font()
    if font_path:
        _PISA_FONT_FAMILY.append("HTUni")
    else:
        _PISA_FONT_FAMILY.append("Helvetica")
    return _PISA_FONT_FAMILY[0]


def _resolve_ht_font(uri: str) -> str | None:
    """Map an ``ht-font://...`` URI to a local TTF path."""
    if not uri.startswith("ht-font://"):
        return None
    regular = _find_unicode_font()
    if not regular:
        return None
    variant = uri[len("ht-font://") :]
    if variant in ("", "regular"):
        return regular
    return _sibling_font(regular, variant) or regular


def _pisa_link_callback(uri: str, rel: str) -> str:
    """``link_callback`` for ``pisa.CreatePDF`` — strict allow-list.

    xhtml2pdf's default fetcher will happily resolve ``file:///etc/passwd``
    or arbitrary ``http://internal/`` URLs from a ``<img src=...>`` /
    ``<link href=...>`` and embed the bytes into the rendered PDF, leaking
    local files or driving SSRF from the operator's host. We block
    everything except:

    * ``ht-font://`` — our internal font scheme (safe; resolved to bundled
      / system TTF paths).
    * ``data:`` — base64-inlined assets the caller already chose to embed
      (no fetch performed by xhtml2pdf).

    For anything else we return an empty string; xhtml2pdf logs a warning
    and renders the asset as broken / missing — the right failure mode for
    a server-side renderer that should never read the local filesystem on
    behalf of attacker-supplied HTML.
    """
    u = (uri or "").strip()
    if not u:
        return ""
    if u.startswith("ht-font://"):
        return _resolve_ht_font(u) or ""
    if u.startswith("data:"):
        return u
    return ""


_PISA_INIT_PATCHED = False


def _patch_pisa_for_local_fonts() -> None:
    """Patch ``pisaContext.__init__`` to pre-register our Unicode TTF.

    Why: xhtml2pdf's normal ``@font-face`` URL fetcher writes to a
    ``NamedTemporaryFile`` which on Windows is opened exclusively, so the
    subsequent reportlab ``TTFont(temp_path)`` fails with "Can't open file".
    Working around the temp-file lock by reopening with a copy is fragile.
    Easier: after every fresh ``pisaContext`` is built, call reportlab's
    ``pdfmetrics.registerFont`` directly with a real, stable filesystem path
    and populate ``self.fontList`` so subsequent ``font-family: HTUni`` CSS
    lookups resolve. This entirely bypasses the URL fetcher.
    Idempotent — applies once per process.
    """
    global _PISA_INIT_PATCHED
    if _PISA_INIT_PATCHED:
        return
    try:
        from xhtml2pdf.context import pisaContext
    except Exception as e:
        logger.debug("xhtml2pdf import failed during font patch: %s", e)
        return

    original_init = pisaContext.__init__

    def patched_init(self: Any, *args: Any, **kwargs: Any) -> None:
        original_init(self, *args, **kwargs)
        # Now ``self`` has empty fontList — pre-load HTUni if a TTF is on disk
        regular = _find_unicode_font()
        if not regular:
            return
        try:
            from reportlab.lib.fonts import addMapping
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont

            registered: dict[tuple[int, int], str] = {}
            variants = {
                (0, 0): regular,
                (1, 0): _sibling_font(regular, "bold") or regular,
                (0, 1): _sibling_font(regular, "italic") or regular,
                (1, 1): _sibling_font(regular, "bolditalic") or regular,
            }
            for (bold, italic), path in variants.items():
                full_name = f"htuni_{bold}{italic}"
                if full_name not in self.fontList:
                    pdfmetrics.registerFont(TTFont(full_name, path))
                addMapping("htuni", bold, italic, full_name)
                registered[(bold, italic)] = full_name
            # Tell xhtml2pdf about it: name + uppercase alias
            self.registerFont("htuni", ["HTUni"])
            for full in registered.values():
                self.fontList[full] = full
        except Exception as e:
            logger.debug("HTUni preload in pisaContext failed: %s", e)

    pisaContext.__init__ = patched_init
    _PISA_INIT_PATCHED = True


def _xhtml2pdf_render(html: str, output: Path) -> None:
    """Pipe ``html`` through xhtml2pdf into ``output`` (raises on failure)."""
    from xhtml2pdf import pisa

    output.parent.mkdir(parents=True, exist_ok=True)
    _patch_pisa_for_local_fonts()
    family = _ensure_pisa_unicode_font()
    if family == "HTUni":
        # Body override — the font itself is registered via the patched
        # ``pisaContext.__init__`` so ``font-family: HTUni`` just resolves.
        font_block = (
            "<style>"
            "body, p, h1, h2, h3, h4, h5, h6, table, td, th, div, span, li, a, "
            "strong, em, b, i, u { font-family: HTUni; }"
            "</style>"
        )
        if "<head>" in html:
            html = html.replace("<head>", "<head>" + font_block, 1)
        else:
            html = font_block + html
    with output.open("wb") as fp:
        pisa_status = pisa.CreatePDF(
            html,
            dest=fp,
            encoding="utf-8",
            link_callback=_pisa_link_callback,
        )
    if pisa_status.err:
        raise ValueError(f"HTML → PDF render başarısız ({pisa_status.err} hata).")


def html_to_pdf(html: str, output: Path) -> None:
    """Render an HTML/CSS string to PDF via xhtml2pdf.

    Supports xhtml2pdf's CSS subset: ``@page`` size + margins, basic font
    families, ``font-size``, ``color``, ``background``, simple table layout.
    No flexbox, no JS, no remote ``<link rel=stylesheet>`` fetching.
    """
    if not html or not html.strip():
        raise ValueError("HTML içeriği boş.")
    _xhtml2pdf_render(html, output)


# Cap on the response body for url_to_pdf — defends against multi-GB streams
# that would otherwise read into memory unchecked. HTML pages rarely exceed
# this; legitimate documents above 50 MB are out of scope for an HTML→PDF tool.
_URL_TO_PDF_MAX_BYTES = 50 * 1024 * 1024
_URL_TO_PDF_MAX_REDIRECTS = 5


class _SSRFGuardedRedirectHandler:
    """``HTTPRedirectHandler`` that re-validates each redirect target.

    ``urllib`` follows 302/301 by default with no opportunity to inspect the
    new URL. Without this guard, an attacker-controlled public URL can return
    ``Location: http://127.0.0.1:8000/admin/...`` (or cloud-metadata IPs) and
    bypass the SSRF check entirely. Subclassing the stdlib handler is the
    clean intercept point — every redirect, internal or external, runs
    ``_assert_public_url`` against the new URL before being followed.
    """

    @staticmethod
    def install() -> Any:
        from urllib.parse import urlparse
        from urllib.request import HTTPRedirectHandler, build_opener

        import core

        class _Guarded(HTTPRedirectHandler):
            def redirect_request(self, req, fp, code, msg, headers, newurl):  # type: ignore[no-untyped-def]
                parsed_new = urlparse(newurl)
                if parsed_new.scheme not in ("http", "https"):
                    raise ValueError(f"Yönlendirme reddedildi: {parsed_new.scheme}://")
                core._assert_public_url(parsed_new)
                return super().redirect_request(req, fp, code, msg, headers, newurl)

        return build_opener(_Guarded())


def url_to_pdf(url: str, output: Path, *, timeout: int = 15) -> None:
    """Fetch a URL with stdlib urllib and render the response body as a PDF.

    Only http(s) is allowed. Charset is detected from the Content-Type header
    when possible, otherwise UTF-8 is assumed. Each redirect target is
    re-validated through ``_assert_public_url`` (see
    ``_SSRFGuardedRedirectHandler``) and the response body is capped at
    ``_URL_TO_PDF_MAX_BYTES``.
    """
    from urllib.parse import urlparse
    from urllib.request import Request

    import core

    if not url or not url.strip():
        raise ValueError("URL boş.")
    parsed = urlparse(url.strip())
    if parsed.scheme not in ("http", "https"):
        raise ValueError("Yalnızca http(s) URL'leri desteklenir.")
    if parsed.username or parsed.password:
        raise ValueError("URL'de kullanıcı/şifre kabul edilmez.")
    core._assert_public_url(parsed)
    opener = _SSRFGuardedRedirectHandler.install()
    req = Request(url, headers={"User-Agent": "ht-pdf-converter/1.x"})
    try:
        with opener.open(req, timeout=timeout) as resp:
            content_type = resp.headers.get("Content-Type", "")
            charset = "utf-8"
            for piece in content_type.split(";"):
                piece = piece.strip().lower()
                if piece.startswith("charset="):
                    charset = piece.split("=", 1)[1].strip() or "utf-8"
                    break
            # Stream-read with cap; bail out as soon as we cross the limit
            # rather than buffering the entire response and then checking.
            chunks: list[bytes] = []
            total = 0
            while True:
                chunk = resp.read(64 * 1024)
                if not chunk:
                    break
                total += len(chunk)
                if total > _URL_TO_PDF_MAX_BYTES:
                    raise ValueError(
                        f"URL yanıtı {_URL_TO_PDF_MAX_BYTES // (1024 * 1024)} MB sınırını aştı."
                    )
                chunks.append(chunk)
            raw = b"".join(chunks)
    except ValueError:
        raise
    except Exception as e:
        raise ValueError(f"URL açılamadı: {sanitize_error(e)}") from e
    try:
        html = raw.decode(charset, errors="replace")
    except LookupError:
        html = raw.decode("utf-8", errors="replace")
    html_to_pdf(html, output)


# ----- .docx → PDF --------------------------------------------------------
def _docx_runs_to_html(runs: Any) -> str:
    """Inline run conversion: <strong>/<em>/<u> + escape HTML chars."""
    import html

    out: list[str] = []
    for run in runs:
        text = html.escape(run.text or "")
        if not text:
            continue
        if run.bold:
            text = f"<strong>{text}</strong>"
        if run.italic:
            text = f"<em>{text}</em>"
        if run.underline:
            text = f"<u>{text}</u>"
        out.append(text)
    return "".join(out)


def _docx_paragraph_html(paragraph: Any) -> str:
    """One docx paragraph → one HTML element (h1-h6 / p)."""
    style_name = (paragraph.style.name or "").lower() if paragraph.style else ""
    body = _docx_runs_to_html(paragraph.runs)
    if not body.strip():
        return "<p>&nbsp;</p>"  # preserve blank lines
    if "heading 1" in style_name or style_name == "title":
        return f"<h1>{body}</h1>"
    if "heading 2" in style_name:
        return f"<h2>{body}</h2>"
    if "heading 3" in style_name:
        return f"<h3>{body}</h3>"
    if "heading 4" in style_name:
        return f"<h4>{body}</h4>"
    if "heading" in style_name:
        return f"<h5>{body}</h5>"
    align_map = {0: "left", 1: "center", 2: "right", 3: "justify"}
    align = align_map.get(paragraph.alignment) if paragraph.alignment is not None else None
    if align:
        return f'<p style="text-align:{align}">{body}</p>'
    return f"<p>{body}</p>"


def _docx_table_html(table: Any) -> str:
    import html

    rows: list[str] = []
    for row in table.rows:
        cells: list[str] = []
        for cell in row.cells:
            cell_text = " ".join(
                _docx_runs_to_html(p.runs) or html.escape(p.text or "") for p in cell.paragraphs
            )
            cells.append(f"<td>{cell_text or '&nbsp;'}</td>")
        rows.append(f"<tr>{''.join(cells)}</tr>")
    return f'<table border="1" cellpadding="4" cellspacing="0">{"".join(rows)}</table>'


def _html_base_css() -> str:
    """Return the project's base CSS with the Unicode-capable font family
    pinned in. Called at render time so the Helvetica fallback is only used
    when no system TTF was found at startup."""
    family = _ensure_pisa_unicode_font()
    return f"""
@page {{ size: A4; margin: 1.8cm; }}
body {{ font-family: {family}; font-size: 11pt; color: #111; }}
h1 {{ font-family: {family}; font-size: 20pt; color: #2F5496; margin-top: 0; }}
h2 {{ font-family: {family}; font-size: 16pt; color: #2F5496; }}
h3 {{ font-family: {family}; font-size: 13pt; color: #1f3a6b; }}
p, li, td, th, span, div {{ font-family: {family}; }}
p {{ margin: 4pt 0; line-height: 1.4; }}
table {{ border-collapse: collapse; margin: 6pt 0; width: 100%; }}
td, th {{ border: 1px solid #888; padding: 4pt; vertical-align: top; }}
th {{ background: #e2e8f0; font-weight: bold; }}
"""


def docx_to_pdf(input_path: Path, output: Path) -> None:
    """Convert a .docx (Word) file to PDF via HTML + xhtml2pdf.

    Preserves: headings (Heading 1–6, Title), bold / italic / underline,
    paragraph alignment, tables. Drops: images, SmartArt, macros, complex
    page layouts, headers/footers (xhtml2pdf prints page numbers via @page).
    """
    if input_path.suffix.lower() != ".docx":
        raise ValueError("Yalnızca .docx (modern Word) destekleniyor — eski .doc desteklenmez.")
    from docx import Document  # python-docx
    from docx.oxml.ns import qn

    try:
        doc = Document(str(input_path))
    except Exception as e:
        raise ValueError(f".docx okunamadı: {sanitize_error(e)}") from e

    # Walk body in document order so paragraphs and tables stay interleaved
    parts: list[str] = []
    body = doc.element.body
    para_iter = iter(doc.paragraphs)
    table_iter = iter(doc.tables)
    para_lookup = {p._element: p for p in doc.paragraphs}
    table_lookup = {t._element: t for t in doc.tables}

    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            p = para_lookup.get(child)
            if p is not None:
                parts.append(_docx_paragraph_html(p))
        elif child.tag == qn("w:tbl"):
            t = table_lookup.get(child)
            if t is not None:
                parts.append(_docx_table_html(t))

    # Fallback if structure walk yielded nothing (older docx variants)
    if not parts:
        for p in para_iter:
            parts.append(_docx_paragraph_html(p))
        for t in table_iter:
            parts.append(_docx_table_html(t))

    title = input_path.stem
    html = (
        "<!doctype html><html><head><meta charset='utf-8'>"
        f"<title>{title}</title>"
        f"<style>{_html_base_css()}</style>"
        f"</head><body>{''.join(parts)}</body></html>"
    )
    _xhtml2pdf_render(html, output)


# ----- .xlsx → PDF --------------------------------------------------------
def xlsx_to_pdf(input_path: Path, output: Path, *, sheet: str | None = None) -> None:
    """Convert a .xlsx workbook to PDF — one HTML table per sheet.

    ``sheet=None`` includes every visible sheet (each starts on its own page);
    otherwise pass a sheet name. Formulas are emitted as their cached value
    (openpyxl's ``data_only=True``). Charts and conditional formatting are
    dropped.
    """
    if input_path.suffix.lower() != ".xlsx":
        raise ValueError("Yalnızca .xlsx (modern Excel) destekleniyor — .xls desteklenmez.")
    import html as _html

    from openpyxl import load_workbook

    try:
        wb = load_workbook(str(input_path), data_only=True, read_only=True)
    except Exception as e:
        raise ValueError(f".xlsx okunamadı: {sanitize_error(e)}") from e

    if sheet is not None and sheet not in wb.sheetnames:
        raise ValueError(f"'{sheet}' sayfası bulunamadı.")

    sections: list[str] = []
    target_sheets = [sheet] if sheet else wb.sheetnames

    for s_name in target_sheets:
        ws = wb[s_name]
        if ws.sheet_state != "visible" and not sheet:
            continue
        rows_html: list[str] = []
        first_row = True
        for row in ws.iter_rows(values_only=True):
            cells = "".join(
                f"<{('th' if first_row else 'td')}>{_html.escape('' if c is None else str(c))}</{'th' if first_row else 'td'}>"
                for c in row
            )
            if cells:
                rows_html.append(f"<tr>{cells}</tr>")
            first_row = False
        if not rows_html:
            continue
        section = (
            f"<h2>{_html.escape(s_name)}</h2><table>{''.join(rows_html)}</table><pdf:nextpage />"
        )
        sections.append(section)
    wb.close()

    if not sections:
        raise ValueError("Çevrilecek görünür sayfa yok.")

    title = input_path.stem
    html_full = (
        "<!doctype html><html><head><meta charset='utf-8'>"
        f"<title>{_html.escape(title)}</title>"
        f"<style>{_html_base_css()}\nh2 {{ page-break-before: always; }} h2:first-of-type {{ page-break-before: avoid; }}</style>"
        f"</head><body>{''.join(sections)}</body></html>"
    )
    _xhtml2pdf_render(html_full, output)
