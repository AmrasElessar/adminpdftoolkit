"""OCR pipeline — runs the readtext loop against each rendered page and
materialises Word/Excel/JPG outputs."""

from __future__ import annotations

import shutil
import zipfile
from pathlib import Path

import core
from core import logger, sanitize_error
from state import ocr_store


def ocr_worker(
    token: str,
    pdf_path: Path,
    target: str,
    job_dir: Path,
    orig_name: str,
) -> None:
    """Background OCR thread for one job. ``target`` ∈ ``{word, excel, jpg}``.

    JPG path skips OCR entirely — it just renders + ZIPs the pages. The
    other targets lazy-import the EasyOCR reader (it pulls in torch on
    demand) and write a Word doc or one-sheet-per-page workbook."""
    try:
        import fitz as _fitz
        stem = Path(orig_name).stem

        if target == "jpg":
            jpg_dir = job_dir / stem
            zoom = 200 / 72
            mat = _fitz.Matrix(zoom, zoom)
            doc = _fitz.open(str(pdf_path))
            total = len(doc)
            ocr_store.update(token, total=total, phase="render")
            jpg_dir.mkdir(parents=True, exist_ok=True)
            files: list[Path] = []
            try:
                for i, page in enumerate(doc):
                    pix = page.get_pixmap(matrix=mat, alpha=False)
                    out = jpg_dir / f"sayfa_{i+1:03d}.jpg"
                    pix.save(str(out), jpg_quality=90)
                    pix = None  # release Pixmap eagerly so 100+ page batches don't bloat
                    files.append(out)
                    ocr_store.update(token, current=i + 1)
            finally:
                doc.close()
            zip_out = job_dir / f"{stem}_OCR.zip"
            with zipfile.ZipFile(zip_out, "w", zipfile.ZIP_DEFLATED) as zf:
                for f in files:
                    zf.write(f, arcname=f"{stem}/{f.name}")
            shutil.rmtree(jpg_dir, ignore_errors=True)
            ocr_store.update(
                token,
                output_path=str(zip_out),
                output_name=zip_out.name,
                media_type="application/zip",
                done=True,
            )
            return

        from pdf_converter import get_ocr_reader
        ocr_store.update(token, phase="loading_model")
        reader = get_ocr_reader()

        doc = _fitz.open(str(pdf_path))
        total = len(doc)
        zoom = 200 / 72
        mat = _fitz.Matrix(zoom, zoom)
        ocr_store.update(token, total=total, phase="ocr")

        pages_text: list[str] = []
        try:
            for i, page in enumerate(doc):
                pix = page.get_pixmap(matrix=mat, alpha=False)
                img_bytes = pix.tobytes("png")
                pix = None  # release Pixmap eagerly between pages
                lines = reader.readtext(img_bytes, detail=0, paragraph=True) or []
                pages_text.append("\n".join(lines))
                ocr_store.update(token, current=i + 1)
        finally:
            doc.close()

        ocr_store.update(token, phase="writing")

        if target == "word":
            from docx import Document as DocxDocument
            d = DocxDocument()
            d.add_heading(f"OCR Sonucu: {orig_name}", level=1)
            for idx, txt in enumerate(pages_text, 1):
                d.add_heading(f"Sayfa {idx}", level=2)
                for para in txt.split("\n"):
                    d.add_paragraph(para)
            out = job_dir / f"{stem}_OCR.docx"
            d.save(str(out))
            media = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        else:  # excel
            from openpyxl import Workbook
            wb = Workbook()
            wb.remove(wb.active)
            for idx, txt in enumerate(pages_text, 1):
                ws = wb.create_sheet(title=f"Sayfa {idx}"[:31])
                for r, line in enumerate(txt.split("\n"), 1):
                    ws.cell(row=r, column=1, value=line)
                ws.column_dimensions["A"].width = 120
            out = job_dir / f"{stem}_OCR.xlsx"
            wb.save(str(out))
            media = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"

        ocr_store.update(
            token,
            output_path=str(out),
            output_name=out.name,
            media_type=media,
            done=True,
        )
    except Exception as e:
        logger.exception("ocr worker failed (token=%s)", token)
        ocr_store.update(token, error=sanitize_error(e), done=True)
    finally:
        snap = ocr_store.snapshot(token)
        if snap:
            core.persist_job_state("ocr", token, snap)
