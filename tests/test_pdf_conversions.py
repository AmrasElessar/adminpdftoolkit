"""Tests for the Section B conversion helpers added in v1.2.

Covers ``image_to_pdf``, ``pdf_to_markdown``, ``pdf_to_csv``, ``docx_to_pdf``,
``xlsx_to_pdf``, ``html_to_pdf`` and their HTTP endpoints. ``url_to_pdf`` is
exercised against an in-process FastAPI test server so we never reach the real
internet.
"""
from __future__ import annotations

import io
import threading
import time
from contextlib import contextmanager
from pathlib import Path

import fitz
import pytest
from fastapi import FastAPI
from fastapi.responses import HTMLResponse
from fastapi.testclient import TestClient

import app
import core


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------
@pytest.fixture
def sample_image(tmp_path: Path) -> Path:
    """Build a 200x150 solid-blue PNG."""
    out = tmp_path / "blue.png"
    pix = fitz.Pixmap(fitz.csRGB, fitz.IRect(0, 0, 200, 150))
    pix.set_rect(pix.irect, (0, 80, 200))
    pix.save(str(out))
    return out


@pytest.fixture
def sample_image_2(tmp_path: Path) -> Path:
    out = tmp_path / "red.jpg"
    pix = fitz.Pixmap(fitz.csRGB, fitz.IRect(0, 0, 100, 100))
    pix.set_rect(pix.irect, (255, 0, 0))
    pix.save(str(out))
    return out


@pytest.fixture
def sample_pdf_with_text(tmp_path: Path) -> Path:
    """A 2-page PDF with a heading-sized line and body text."""
    out = tmp_path / "text.pdf"
    doc = fitz.open()
    p1 = doc.new_page()
    p1.insert_text((72, 80), "BAS LIK", fontsize=24, fontname="helv")
    p1.insert_text((72, 130), "Bu bir paragraf metnidir.",
                   fontsize=11, fontname="helv")
    p2 = doc.new_page()
    p2.insert_text((72, 80), "Ikinci sayfa", fontsize=11, fontname="helv")
    doc.save(str(out))
    doc.close()
    return out


@pytest.fixture
def sample_pdf_with_table(tmp_path: Path) -> Path:
    """Tiny PDF embedding a 2x3 table that pdfplumber can extract."""
    out = tmp_path / "table.pdf"
    doc = fitz.open()
    page = doc.new_page()
    # Use lines + insert_textbox to draw a real table grid
    cells = [
        ["Ad", "Soyad", "Yas"],
        ["Ali", "Yilmaz", "30"],
        ["Veli", "Kaya", "45"],
    ]
    x0 = 72.0
    y0 = 100.0
    cw = 90.0
    rh = 24.0
    for r, row in enumerate(cells):
        for c, value in enumerate(row):
            rect = fitz.Rect(x0 + c * cw, y0 + r * rh,
                             x0 + (c + 1) * cw, y0 + (r + 1) * rh)
            page.draw_rect(rect, color=(0, 0, 0), width=0.5)
            page.insert_textbox(rect, value, fontsize=10, fontname="helv",
                                align=fitz.TEXT_ALIGN_CENTER)
    doc.save(str(out))
    doc.close()
    return out


@pytest.fixture
def sample_docx(tmp_path: Path) -> Path:
    """Build a small .docx with a heading, bold text, and a 2x2 table."""
    from docx import Document

    out = tmp_path / "sample.docx"
    d = Document()
    d.add_heading("Test Belgesi", level=1)
    p = d.add_paragraph()
    p.add_run("Bu cümle ").italic = False
    p.add_run("kalın").bold = True
    p.add_run(" yazıyla bitiyor.")
    d.add_paragraph("İkinci paragraf — düz metin.")
    table = d.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "A"
    table.rows[0].cells[1].text = "B"
    table.rows[1].cells[0].text = "1"
    table.rows[1].cells[1].text = "2"
    d.save(str(out))
    return out


@pytest.fixture
def sample_xlsx(tmp_path: Path) -> Path:
    from openpyxl import Workbook

    out = tmp_path / "sample.xlsx"
    wb = Workbook()
    ws = wb.active
    ws.title = "Veri"
    ws.append(["Ad", "Soyad", "Yas"])
    ws.append(["Ali", "Yilmaz", 30])
    ws.append(["Veli", "Kaya", 45])
    ws2 = wb.create_sheet("Diğer")
    ws2.append(["x", "y"])
    ws2.append([1, 2])
    wb.save(str(out))
    return out


@pytest.fixture
def client(monkeypatch: pytest.MonkeyPatch) -> TestClient:
    monkeypatch.setattr(core, "is_local_request", lambda req: True)
    monkeypatch.setattr("core.is_local_request", lambda req: True)
    return TestClient(app.app)


# ---------------------------------------------------------------------------
# core.image_to_pdf
# ---------------------------------------------------------------------------
def test_image_to_pdf_concatenates(sample_image: Path, sample_image_2: Path, tmp_path: Path) -> None:
    out = tmp_path / "imgs.pdf"
    n = core.image_to_pdf([sample_image, sample_image_2], out)
    assert n == 2
    with fitz.open(str(out)) as d:
        assert d.page_count == 2


def test_image_to_pdf_empty_raises(tmp_path: Path) -> None:
    with pytest.raises(ValueError):
        core.image_to_pdf([], tmp_path / "x.pdf")


def test_image_to_pdf_unknown_ext_raises(tmp_path: Path) -> None:
    bad = tmp_path / "fake.txt"
    bad.write_text("not an image")
    with pytest.raises(ValueError):
        core.image_to_pdf([bad], tmp_path / "x.pdf")


# ---------------------------------------------------------------------------
# core.pdf_to_markdown
# ---------------------------------------------------------------------------
def test_pdf_to_markdown_emits_text(sample_pdf_with_text: Path, tmp_path: Path) -> None:
    out = tmp_path / "out.md"
    pages = core.pdf_to_markdown(sample_pdf_with_text, out)
    assert pages == 2
    md = out.read_text(encoding="utf-8")
    assert "Bu bir paragraf metnidir" in md
    # The 24-pt line should be tagged as a heading (any level)
    assert any(line.lstrip().startswith("#") for line in md.splitlines())


# ---------------------------------------------------------------------------
# core.pdf_to_csv
# ---------------------------------------------------------------------------
def test_pdf_to_csv_extracts_table(sample_pdf_with_table: Path, tmp_path: Path) -> None:
    out = tmp_path / "out.csv"
    rows = core.pdf_to_csv(sample_pdf_with_table, out)
    assert rows >= 3  # at least header + 2 rows
    body = out.read_text(encoding="utf-8-sig")
    assert "Ad" in body
    assert "Soyad" in body
    assert "Ali" in body


def test_pdf_to_csv_no_table_raises(sample_pdf_with_text: Path, tmp_path: Path) -> None:
    with pytest.raises(ValueError):
        core.pdf_to_csv(sample_pdf_with_text, tmp_path / "out.csv")


# ---------------------------------------------------------------------------
# core.html_to_pdf
# ---------------------------------------------------------------------------
def test_html_to_pdf_basic(tmp_path: Path) -> None:
    out = tmp_path / "out.pdf"
    core.html_to_pdf("<h1>Merhaba</h1><p>Dünya</p>", out)
    assert out.exists()
    with fitz.open(str(out)) as d:
        assert d.page_count >= 1
        text = "\n".join(p.get_text() for p in d)
    assert "Merhaba" in text


def test_html_to_pdf_empty_raises(tmp_path: Path) -> None:
    with pytest.raises(ValueError):
        core.html_to_pdf("", tmp_path / "x.pdf")


# ---------------------------------------------------------------------------
# core.docx_to_pdf
# ---------------------------------------------------------------------------
def test_docx_to_pdf(sample_docx: Path, tmp_path: Path) -> None:
    out = tmp_path / "out.pdf"
    core.docx_to_pdf(sample_docx, out)
    assert out.exists()
    with fitz.open(str(out)) as d:
        text = "\n".join(p.get_text() for p in d)
    assert "Test Belgesi" in text
    assert "İkinci paragraf" in text


def test_docx_to_pdf_rejects_doc(tmp_path: Path) -> None:
    fake = tmp_path / "old.doc"
    fake.write_bytes(b"\xd0\xcf\x11\xe0")  # OLE header
    with pytest.raises(ValueError):
        core.docx_to_pdf(fake, tmp_path / "x.pdf")


# ---------------------------------------------------------------------------
# core.xlsx_to_pdf
# ---------------------------------------------------------------------------
def test_xlsx_to_pdf(sample_xlsx: Path, tmp_path: Path) -> None:
    out = tmp_path / "out.pdf"
    core.xlsx_to_pdf(sample_xlsx, out)
    assert out.exists()
    with fitz.open(str(out)) as d:
        text = "\n".join(p.get_text() for p in d)
    assert "Veri" in text
    assert "Ali" in text
    assert "Diğer" in text  # second sheet present


def test_xlsx_to_pdf_specific_sheet(sample_xlsx: Path, tmp_path: Path) -> None:
    out = tmp_path / "out.pdf"
    core.xlsx_to_pdf(sample_xlsx, out, sheet="Diğer")
    with fitz.open(str(out)) as d:
        text = "\n".join(p.get_text() for p in d)
    # Only "Diğer" sheet should be rendered
    assert "Diğer" in text
    assert "Ali" not in text  # came from "Veri"


def test_xlsx_to_pdf_unknown_sheet_raises(sample_xlsx: Path, tmp_path: Path) -> None:
    with pytest.raises(ValueError):
        core.xlsx_to_pdf(sample_xlsx, tmp_path / "x.pdf", sheet="Yok")


# ---------------------------------------------------------------------------
# core.url_to_pdf — point at a local FastAPI server fixture
# ---------------------------------------------------------------------------
@contextmanager
def _local_html_server(html: str):
    """Spin up a tiny HTTP server returning ``html`` on /, on a random port."""
    import socket
    import uvicorn

    sub_app = FastAPI()

    @sub_app.get("/", response_class=HTMLResponse)
    def index() -> str:
        return html

    sock = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
    sock.bind(("127.0.0.1", 0))
    port = sock.getsockname()[1]
    sock.close()

    config = uvicorn.Config(sub_app, host="127.0.0.1", port=port,
                            log_level="error", access_log=False)
    server = uvicorn.Server(config)
    thread = threading.Thread(target=server.run, daemon=True)
    thread.start()
    # Wait for the socket to come up
    for _ in range(50):
        try:
            s = socket.create_connection(("127.0.0.1", port), timeout=0.2)
            s.close()
            break
        except OSError:
            time.sleep(0.05)
    try:
        yield f"http://127.0.0.1:{port}/"
    finally:
        server.should_exit = True
        thread.join(timeout=3)


def test_url_to_pdf(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    # SSRF guard would block 127.0.0.1; this is a controlled integration test
    # so we bypass the public-URL assertion just for this case.
    monkeypatch.setattr(core, "_assert_public_url", lambda parsed: None)
    out = tmp_path / "url.pdf"
    with _local_html_server("<h1>From URL</h1><p>Test</p>") as url:
        core.url_to_pdf(url, out, timeout=5)
    assert out.exists()
    with fitz.open(str(out)) as d:
        text = "\n".join(p.get_text() for p in d)
    assert "From URL" in text


def test_url_to_pdf_rejects_non_http(tmp_path: Path) -> None:
    with pytest.raises(ValueError):
        core.url_to_pdf("file:///etc/passwd", tmp_path / "x.pdf")


# ---------------------------------------------------------------------------
# Endpoint smoke tests
# ---------------------------------------------------------------------------
def test_endpoint_from_images(client: TestClient, sample_image: Path, sample_image_2: Path) -> None:
    r = client.post(
        "/pdf/from-images",
        files=[
            ("files", (sample_image.name, sample_image.read_bytes(), "image/png")),
            ("files", (sample_image_2.name, sample_image_2.read_bytes(), "image/jpeg")),
        ],
        data={"output_name": "iki-resim"},
    )
    assert r.status_code == 200, r.text
    assert r.headers["content-type"] == "application/pdf"
    with fitz.open(stream=r.content, filetype="pdf") as d:
        assert d.page_count == 2


def test_endpoint_from_images_rejects_non_image(client: TestClient) -> None:
    r = client.post(
        "/pdf/from-images",
        files=[("files", ("a.txt", b"hello", "text/plain"))],
    )
    assert r.status_code == 400


def test_endpoint_to_markdown(client: TestClient, sample_pdf_with_text: Path) -> None:
    r = client.post(
        "/pdf/to-markdown",
        files={"file": ("text.pdf", sample_pdf_with_text.read_bytes(), "application/pdf")},
    )
    assert r.status_code == 200, r.text
    assert "markdown" in r.headers["content-type"]
    assert b"paragraf" in r.content.lower() or b"BAS" in r.content


def test_endpoint_to_csv(client: TestClient, sample_pdf_with_table: Path) -> None:
    r = client.post(
        "/pdf/to-csv",
        files={"file": ("t.pdf", sample_pdf_with_table.read_bytes(), "application/pdf")},
        data={"delimiter": ","},
    )
    assert r.status_code == 200, r.text
    assert "csv" in r.headers["content-type"]
    text = r.content.decode("utf-8-sig")
    assert "Ad" in text
    assert "Ali" in text


def test_endpoint_to_csv_no_table_400(client: TestClient, sample_pdf_with_text: Path) -> None:
    r = client.post(
        "/pdf/to-csv",
        files={"file": ("t.pdf", sample_pdf_with_text.read_bytes(), "application/pdf")},
    )
    assert r.status_code == 400


def test_endpoint_from_docx(client: TestClient, sample_docx: Path) -> None:
    r = client.post(
        "/pdf/from-docx",
        files={"file": ("doc.docx", sample_docx.read_bytes(),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    assert r.status_code == 200, r.text
    assert r.headers["content-type"] == "application/pdf"
    with fitz.open(stream=r.content, filetype="pdf") as d:
        text = "\n".join(p.get_text() for p in d)
    assert "Test Belgesi" in text


def test_endpoint_from_docx_rejects_old_doc(client: TestClient) -> None:
    r = client.post(
        "/pdf/from-docx",
        files={"file": ("old.doc", b"\xd0\xcf\x11\xe0", "application/msword")},
    )
    assert r.status_code == 400


def test_endpoint_from_xlsx(client: TestClient, sample_xlsx: Path) -> None:
    r = client.post(
        "/pdf/from-xlsx",
        files={"file": ("d.xlsx", sample_xlsx.read_bytes(),
                        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")},
    )
    assert r.status_code == 200, r.text
    assert r.headers["content-type"] == "application/pdf"
    with fitz.open(stream=r.content, filetype="pdf") as d:
        text = "\n".join(p.get_text() for p in d)
    assert "Veri" in text


def test_endpoint_from_html(client: TestClient) -> None:
    r = client.post(
        "/pdf/from-html",
        data={"html": "<h1>Başlık</h1><p>İçerik</p>", "output_name": "test"},
    )
    assert r.status_code == 200, r.text
    assert r.headers["content-type"] == "application/pdf"
    with fitz.open(stream=r.content, filetype="pdf") as d:
        text = "\n".join(p.get_text() for p in d)
    assert "Başlık" in text


def test_endpoint_from_html_empty_400(client: TestClient) -> None:
    r = client.post("/pdf/from-html", data={"html": "   "})
    assert r.status_code == 400


def test_endpoint_from_url_invalid_scheme_400(client: TestClient) -> None:
    r = client.post("/pdf/from-url", data={"url": "ftp://x"})
    assert r.status_code == 400
